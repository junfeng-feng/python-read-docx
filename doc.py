# coding=utf-8
#
# soffice --headless --convert-to docx:"Office Open XML Text" --outdir=/home/ *.doc 
# soffice have bug in convert-to docx
import docx
import os
import traceback
from pip._vendor.distlib._backport.tarfile import TUREAD
from _ast import If

writeFileFlag = True
testFlag = True
directoryName = "/Users/desktop/Downloads/0116"  # all Data
#directoryName = "/Users/desktop/Downloads/1230_parsed" 
#directoryName = "/Users/desktop/Downloads/docx_test"

selectedList = ["00025175", "009E51D1", "F052", "0052", "F0FE", "F0A2", "003C3A24", "007752F5", "001C42A3", "000A5B43"]
class Element():
    def __init__(self):
        self.xml = ""
class Paragrah():
    def __init__(self):
        self.text = ""
        self._element = Element()

def replaceText(text):
    return text.replace(" ", "").replace("（", "").replace("）",
            "").replace("姓名", "").replace("：", "").replace(":",
            "").replace("√", "").replace("_",
            "").replace("_", "").replace("其他请补充填写", ""
            ).replace('调查员',
        "").replace("请调查员记录", "").replace("使用其他方式查灾请补充填写",
    "").replace("请记录", "").replace("；", "").replace("✅", "").replace("请补充填写", "")


# 取基本信息，括号内的内容
def getTextByPosition(data, rowno, colno, pno):
    try:
        return getTextin(data[rowno][colno][pno].text)
    except Exception as e:
        return ""
    pass

def getParagrahByPostion(data, rowno, colno, pno):
    try:
#         if pno ==48:
#             print(data[rowno][colno][pno]._element.xml)
        return data[rowno][colno][pno]
    except Exception as e:
        return Paragrah()
    pass

def getTextin(text):
    if text.find("括号内填写原信息，下同") != -1:
        return replaceText(text)
    try:
        # 取括号内的值
        # （可以考虑直接全部采用替换，讲关键字全部替换掉即可）
        return replaceText(text.split("：")[1])
    except Exception as e:
        open("error.log", "a").write(str(e) + "\n")
        return replaceText(text)

    
def isSelected(paragrahObj):
    global selectedList
    #特殊的字符 √
    if paragrahObj._element.xml.find("√")!=-1:
        return True
    
    for select in selectedList:
        if paragrahObj._element.xml.find('''"'''+select+'''"''') != -1:
            #print(paragrahObj._element.xml + "\n"+select+"12323")
            return True
        
    # 纯方框没选中
    if paragrahObj._element.xml.find('''F0FE''') == -1:
        # 未选中
        return False

def multiSelect(pList):
    text = ""
    for p in pList:
        if isSelected(p):
            text += replaceText(p.text) + "，"
    return text
    pass

def simpleSelect(paragrahObj):
    if isSelected(paragrahObj):
        return "是"
    else:
        return "否"

def getUpdate(paragrahObj):
    text = paragrahObj._element.xml.replace("灾害信息员是否有变更", "")
    yesIndex = text.find("是")
    noIndex = text.find("否")
    
    global selectedList
    for select in selectedList:
        selectIndex = text.find(select)
        if selectIndex != -1:
            if selectIndex < yesIndex:
                return "是"
            else:
                return "否"
            
    selectedListLocal = ["√"]
    for select in selectedListLocal:
        selectIndex = text.find(select)
        if selectIndex != -1 :
            if selectIndex < noIndex:
                return "是"
            else:
                return "否"
    print("last", yesIndex, noIndex, selectIndex)    
    return ""


def parseFile(fileReult, filename):
    global writeFileFlag
    doc = docx.Document(filename) 
    table = doc.tables[0]
    data = {}
    yangqiaohuiFlag = False
    for rowno, row in enumerate(table.rows):
        rowCellsText = []
        for colno, cell in enumerate(row.cells):
            if not cell.text in rowCellsText:
                rowCellsText.append(cell.text) 
                # print("rowno %s  colno %s: %s"%(rowno, colno, cell.text))
                for pno, p in enumerate(cell.paragraphs):
                    if writeFileFlag:
                        open("text.txt", "a").write(filename + "\n")
                        open("text.txt", "a").write("rowno %s, colno %s, pno:%s, %s, %s\n" % (rowno, colno, pno, p.text, ''))
                    if not rowno in data:
                        data[rowno] = {}
                    if not colno in data[rowno]:
                        data[rowno][colno] = {}
                    if not pno in data[rowno][colno]:
                        data[rowno][colno][pno] = {}
                    if p._element.xml.find("灾害信息员工作情况")!=-1 and rowno ==10:
                        yangqiaohuiFlag = True
                        
                    if yangqiaohuiFlag and rowno >=10 and rowno < 18:
                        #特殊处理
                        data[rowno -1][colno][pno] = p
                    else:
                        data[rowno][colno][pno] = p
            
    base8Pno = 0
    if getParagrahByPostion(data,8,0,21).text.find("如为兼职，每天用于灾害信息员的工作时长") == -1:
        # 缺少两个问题：1.如为兼职， 2.兼职时长，
        base8Pno = -8
        pass
    print ("base8Pno:" +  str(base8Pno))

    fileResult.write(",")  # 序号
    fileResult.write(getUpdate(data[1][0][0]) + ",")  # 灾害信息员是否有变更
    fileResult.write(getTextByPosition(data, 2, 0, 0).replace("括号内填写原信息，下同", "") + ",")  # 姓名
    fileResult.write(getTextByPosition(data, 2, 1, 0) + ",")  # 性别
    fileResult.write(getTextByPosition(data, 2, 2, 0) + ",")  # 年龄
    fileResult.write(getTextByPosition(data, 3, 0, 0) + ",")  # 学历
    fileResult.write(getTextByPosition(data, 3, 1, 0) + ",")  # 专业
    fileResult.write(getTextByPosition(data, 3, 2, 0) + ",")  # 职称
    fileResult.write(getTextByPosition(data, 4, 0, 0) + ",")  # 政治面貌
    fileResult.write(getTextByPosition(data, 4, 1, 0).replace("座机", "") + ",")  # 移动电话
    fileResult.write(getTextByPosition(data, 4, 1, 1) + ",")  # 座机
    fileResult.write(getTextByPosition(data, 6, 0, 0) + ",")  # 区县
    fileResult.write(getTextByPosition(data, 6, 1, 0) + ",")  # 乡镇（街道）
    fileResult.write(getTextByPosition(data, 6, 2, 0) + ",")  # 村（社区）
    fileResult.write(getTextByPosition(data, 7, 0, 0) + ",")  # 单位
    fileResult.write(getTextByPosition(data, 7, 1, 0) + ",")  # 工龄
    fileResult.write(getTextByPosition(data, 7, 2, 0) + ",")  # 职务
    fileResult.write(multiSelect([getParagrahByPostion(data,8,0,2), getParagrahByPostion(data,8,0,3)]) + ",")  # 是否为原民政灾害信息员
    fileResult.write(multiSelect([getParagrahByPostion(data,8,0,7),
          getParagrahByPostion(data,8,0,6),]));fileResult.write(",")  # 是否为地震速报员
    fileResult.write(simpleSelect(getParagrahByPostion(data,8,0,11)) + ",")  # 是否为安全员
    fileResult.write(multiSelect([getParagrahByPostion(data,8,0,14),
                     getParagrahByPostion(data,8,0,15),]));fileResult.write(",")  # 如是，继续提问
    
    fileResult.write(simpleSelect(getParagrahByPostion(data,8,0,17)));fileResult.write(",")  
    fileResult.write(multiSelect([
        getParagrahByPostion(data,8,0,18),
        getParagrahByPostion(data,8,0,19)]));fileResult.write(",")  # 是否为专职灾害信息员
    #fileResult.write(
    #         replaceText(getParagrahByPostion(data,8,0,14).text)  if(not isSelected(getParagrahByPostion(data,8,0,14))) else  "");fileResult.write(",")  
    fileResult.write(multiSelect([getParagrahByPostion(data,8,0,base8Pno + 24),
              getParagrahByPostion(data,8,0,base8Pno + 23),
              getParagrahByPostion(data,8,0,base8Pno + 22),])) ;fileResult.write(",")  # 如为兼职，每天用于灾害信息员的工作时长
    fileResult.write(replaceText(getParagrahByPostion(data,8,0,base8Pno + 27).text.replace("否，还兼职其他工作", "")) + ","),  # 如为兼职，您的本职工作是
    fileResult.write(multiSelect([getParagrahByPostion(data,8,0,base8Pno + 33),
        getParagrahByPostion(data,8,0,base8Pno + 32),
         getParagrahByPostion(data,8,0,base8Pno + 31),
         getParagrahByPostion(data,8,0,base8Pno + 30),
        getParagrahByPostion(data,8,0,base8Pno + 34),])) ;fileResult.write(",")  # 担任灾害信息员的工作年限
    fileResult.write("灾害信息员培训证书" if isSelected(getParagrahByPostion(data,8,0,base8Pno + 38)) else "")    ;fileResult.write(",")   
    fileResult.write("安全员" if isSelected(getParagrahByPostion(data,8,0,base8Pno + 39)) else "")      ;fileResult.write(",") 
    fileResult.write("注册安全工程师" if isSelected(getParagrahByPostion(data,8,0,base8Pno + 40)) else "")  ;fileResult.write(",") 
    fileResult.write(replaceText(getParagrahByPostion(data,8,0,41).text) if isSelected(getParagrahByPostion(data,8,0,base8Pno + 41)) else "")  ;fileResult.write(",")     
    fileResult.write("以上皆无" if isSelected(getParagrahByPostion(data,8,0,base8Pno + 42)) else "")  ;fileResult.write(",") 
    fileResult.write(multiSelect([
           getParagrahByPostion(data,8,0,base8Pno + 46),
            getParagrahByPostion(data,8,0,base8Pno + 47),
            getParagrahByPostion(data,8,0,base8Pno + 48),
            ])) ;fileResult.write(",")  # 每年是否安排体检
#     print(multiSelect([
#            getParagrahByPostion(data,8,0,base8Pno + 46),
#             getParagrahByPostion(data,8,0,base8Pno + 47),
#             getParagrahByPostion(data,8,0,base8Pno + 48),
#             ]))
    fileResult.write(multiSelect([getParagrahByPostion(data,8,0,base8Pno + 51),
             getParagrahByPostion(data,8,0,base8Pno + 52),
             getParagrahByPostion(data,8,0,base8Pno + 53),
             getParagrahByPostion(data,8,0,base8Pno + 54),])) ;fileResult.write(",")  # 您认为，目前自己的健康状况如何
    #print(isSelected(getParagrahByPostion(data,8,0,base8Pno + 52)))
    # 请问您的健康状况是否在以下方面存在问题
    fileResult.write(replaceText(getParagrahByPostion(data,8,0,base8Pno + 57).text) if isSelected(getParagrahByPostion(data,8,0,base8Pno + 57)) else "") ;fileResult.write(",") 
    fileResult.write(replaceText(getParagrahByPostion(data,8,0,base8Pno + 58).text) if isSelected(getParagrahByPostion(data,8,0,base8Pno + 58)) else "") ;fileResult.write(",") 
    fileResult.write(replaceText(getParagrahByPostion(data,8,0,base8Pno + 59).text) if isSelected(getParagrahByPostion(data,8,0,base8Pno + 59)) else "") ;fileResult.write(",") 
    fileResult.write(replaceText(getParagrahByPostion(data,8,0,base8Pno + 60).text) if isSelected(getParagrahByPostion(data,8,0,base8Pno + 60)) else "") ;fileResult.write(",") 
    fileResult.write(replaceText(getParagrahByPostion(data,8,0,base8Pno + 61).text) if isSelected(getParagrahByPostion(data,8,0,base8Pno + 61)) else "") ;fileResult.write(",") 
    fileResult.write(replaceText(getParagrahByPostion(data,8,0,base8Pno + 62).text) if isSelected(getParagrahByPostion(data,8,0,base8Pno + 62)) else "") ;fileResult.write(",") 
    fileResult.write(replaceText(getParagrahByPostion(data,8,0,base8Pno + 63).text) if isSelected(getParagrahByPostion(data,8,0,base8Pno + 63)) else "") ;fileResult.write(",") 
    fileResult.write(replaceText(getParagrahByPostion(data,8,0,base8Pno + 64).text) if isSelected(getParagrahByPostion(data,8,0,base8Pno + 64)) else "") ;fileResult.write(",") 
    fileResult.write(replaceText(getParagrahByPostion(data,8,0,base8Pno + 65).text) if isSelected(getParagrahByPostion(data,8,0,base8Pno + 65)) else "") ;fileResult.write(",")  
    
    fileResult.write(multiSelect([getParagrahByPostion(data,8,0,base8Pno + 70),
                                  getParagrahByPostion(data,8,0,base8Pno + 69)]) 
        ) ;fileResult.write(",")  # 您对目前从事的灾害信息员工作是否满意
    fileResult.write(getParagrahByPostion(data,8,0,base8Pno + 72).text + getParagrahByPostion(data,8,0,base8Pno + 73).text + ",")  # 存在问题
    fileResult.write(getParagrahByPostion(data,8,0,base8Pno + 76).text + ",")  # 建议
    
    fileResult.write("手机" if(isSelected(getParagrahByPostion(data,10,0,1))) else "");fileResult.write(",")
    fileResult.write("座机" if(isSelected(getParagrahByPostion(data,10,0,2))) else "");fileResult.write(",")
    fileResult.write("传真" if(isSelected(getParagrahByPostion(data,10,0,3))) else "");fileResult.write(",")
    fileResult.write("电子邮件" if(isSelected(getParagrahByPostion(data,10,0,4))) else "");fileResult.write(",")
    fileResult.write(replaceText(getParagrahByPostion(data,10,0,5).text) if(isSelected(getParagrahByPostion(data,10,0,1))) else "");fileResult.write(",")
    fileResult.write("在村长、村支书、村干部带领下查灾，需要参考领导的意见" if(isSelected(getParagrahByPostion(data,10,0,8))) else "");fileResult.write(",")
    fileResult.write("独立查灾" if(isSelected(getParagrahByPostion(data,10,0,9))) else "");fileResult.write(",")
    fileResult.write(replaceText(getParagrahByPostion(data,10,0,10).text) if(not isSelected(getParagrahByPostion(data,10,0,10))) else  "");fileResult.write(",")
    # 影响您不能及时报灾的因素有
    fileResult.write(getParagrahByPostion(data,10,0,13).text if (isSelected(getParagrahByPostion(data,10,0,13))) else "");fileResult.write(",")
    fileResult.write(getParagrahByPostion(data,10,0,14).text if (isSelected(getParagrahByPostion(data,10,0,14))) else "");fileResult.write(",")
    fileResult.write(getParagrahByPostion(data,10,0,15).text if (isSelected(getParagrahByPostion(data,10,0,15))) else "");fileResult.write(",")
    fileResult.write(getParagrahByPostion(data,10,0,16).text if (isSelected(getParagrahByPostion(data,10,0,16))) else "");fileResult.write(",")
    fileResult.write(getParagrahByPostion(data,10,0,17).text if (isSelected(getParagrahByPostion(data,10,0,17))) else "");fileResult.write(",")
    fileResult.write(getParagrahByPostion(data,10,0,18).text if (isSelected(getParagrahByPostion(data,10,0,18))) else "");fileResult.write(",")
    fileResult.write(getParagrahByPostion(data,10,0,19).text if (isSelected(getParagrahByPostion(data,10,0,19))) else "");fileResult.write(",")
    fileResult.write(getParagrahByPostion(data,10,0,20).text if (isSelected(getParagrahByPostion(data,10,0,20))) else "");fileResult.write(",")
    fileResult.write(replaceText(getParagrahByPostion(data,10,0,21).text) if (isSelected(getParagrahByPostion(data,10,0,21))) else "");fileResult.write(",")                
                    
    fileResult.write(multiSelect([getParagrahByPostion(data,10,0,23),getParagrahByPostion(data,10,0,24)]) + ",")  # 是否上报过自然灾害灾情信息
    fileResult.write(replaceText(getParagrahByPostion(data,10,0,27).text) + ",")  #  报灾过程
    fileResult.write(replaceText(getParagrahByPostion(data,10,0,35).text + getParagrahByPostion(data,10,0,36).text) + ",")  # 救灾过程
    fileResult.write(multiSelect([getParagrahByPostion(data,12,0,1),getParagrahByPostion(data,12,0,2)]) + ",") ,  # 是否接受过灾害信息员的相关培训
    fileResult.write(multiSelect([
                   getParagrahByPostion(data,12,0,5),
                   getParagrahByPostion(data,12,0,6),
                   getParagrahByPostion(data,12,0,7),
                   getParagrahByPostion(data,12,0,8),
                   getParagrahByPostion(data,12,0,9)]
                   )
                    + ",")  # 最近一次接受灾害信息员相关培训是在
    fileResult.write(multiSelect([
                   getParagrahByPostion(data,12,0,12),
                   getParagrahByPostion(data,12,0,13),
                   getParagrahByPostion(data,12,0,14),
                   getParagrahByPostion(data,12,0,15),
                   getParagrahByPostion(data,12,0,16)
                   ]
                   ) + ",")  # 一共接受过多少次灾害信息员相关培训
    fileResult.write(multiSelect([
                   getParagrahByPostion(data,12,0,19),
                   getParagrahByPostion(data,12,0,20),
                   getParagrahByPostion(data,12,0,21),
                   ]
                   ) + ",")  # 培训的课程内容是否实用
    fileResult.write(multiSelect([
                   getParagrahByPostion(data,12,0,24),
                   getParagrahByPostion(data,12,0,25),
                   getParagrahByPostion(data,12,0,26),
                   getParagrahByPostion(data,12,0,27),
                   ]
                   ) + ",")  # 如果不实用，您认为主要原因是
    fileResult.write(replaceText(getParagrahByPostion(data,12,0,30).text + getParagrahByPostion(data,12,0,31).text + getParagrahByPostion(data,12,0,32).text) + ",")  # 您最希望得到哪方面的培训
    fileResult.write(getParagrahByPostion(data,12,0,35).text + ",") ,  # 如对培训工作有其他意见或建议，请说明
    fileResult.write(multiSelect([
                   getParagrahByPostion(data,14,0,1),
                   getParagrahByPostion(data,14,0,2),
                   getParagrahByPostion(data,14,0,3),
                   ]
                   ) + ",")  # 作为灾害信息员的收入或补贴（月薪）情况
    fileResult.write(multiSelect([
                   getParagrahByPostion(data,14,0,6),
                   getParagrahByPostion(data,14,0,7),
                   getParagrahByPostion(data,14,0,8),
                   ]
                   ) + ",")  # 是否有灾害信息员工作意外伤害保险
    fileResult.write(multiSelect([
                   getParagrahByPostion(data,14,0,11),
                   getParagrahByPostion(data,14,0,12),
                   getParagrahByPostion(data,14,0,13),
                   getParagrahByPostion(data,14,0,14)  
                   ]
        ) + ",")  # 是否有相应的奖励制度及表彰制度 
    # 是否发放过个人防护用品（多选1）
    fileResult.write(getParagrahByPostion(data,14,0,17).text if (isSelected(getParagrahByPostion(data,14,0,17))) else "");fileResult.write(",")
    fileResult.write(getParagrahByPostion(data,14,0,18).text if (isSelected(getParagrahByPostion(data,14,0,18))) else "");fileResult.write(",")
    fileResult.write(getParagrahByPostion(data,14,0,19).text if (isSelected(getParagrahByPostion(data,14,0,19))) else "");fileResult.write(",")
    fileResult.write(getParagrahByPostion(data,14,0,20).text if (isSelected(getParagrahByPostion(data,14,0,20))) else "");fileResult.write(",")
    fileResult.write(getParagrahByPostion(data,14,0,21).text if (isSelected(getParagrahByPostion(data,14,0,21))) else "");fileResult.write(",")
    fileResult.write(getParagrahByPostion(data,14,0,22).text if (isSelected(getParagrahByPostion(data,14,0,22))) else "");fileResult.write(",")
    fileResult.write(getParagrahByPostion(data,14,0,23).text if (isSelected(getParagrahByPostion(data,14,0,23))) else "");fileResult.write(",")
    fileResult.write(getParagrahByPostion(data,14,0,24).text if (isSelected(getParagrahByPostion(data,14,0,24))) else "");fileResult.write(",")
    fileResult.write(getParagrahByPostion(data,14,0,25).text if (isSelected(getParagrahByPostion(data,14,0,25))) else "");fileResult.write(",")
    fileResult.write(getParagrahByPostion(data,14,0,26).text if (isSelected(getParagrahByPostion(data,14,0,26))) else "");fileResult.write(",")
    fileResult.write(getParagrahByPostion(data,14,0,27).text if (isSelected(getParagrahByPostion(data,14,0,27))) else "");fileResult.write(",")  # 其他
    # 作为灾害信息员，您最急需的个人防护用品是（多选）
    fileResult.write(getParagrahByPostion(data,14,0,29).text if (isSelected(getParagrahByPostion(data,14,0,29))) else "");fileResult.write(",")
    fileResult.write(getParagrahByPostion(data,14,0,30).text if (isSelected(getParagrahByPostion(data,14,0,30))) else "");fileResult.write(",")
    fileResult.write(getParagrahByPostion(data,14,0,31).text if (isSelected(getParagrahByPostion(data,14,0,31))) else "");fileResult.write(",")
    fileResult.write(getParagrahByPostion(data,14,0,32).text if (isSelected(getParagrahByPostion(data,14,0,32))) else "");fileResult.write(",")
    fileResult.write(getParagrahByPostion(data,14,0,33).text if (isSelected(getParagrahByPostion(data,14,0,33))) else "");fileResult.write(",")
    fileResult.write(getParagrahByPostion(data,14,0,34).text if (isSelected(getParagrahByPostion(data,14,0,34))) else "");fileResult.write(",")
    fileResult.write(getParagrahByPostion(data,14,0,35).text if (isSelected(getParagrahByPostion(data,14,0,35))) else "");fileResult.write(",")
    fileResult.write(getParagrahByPostion(data,14,0,36).text if (isSelected(getParagrahByPostion(data,14,0,36))) else "");fileResult.write(",")
    fileResult.write(replaceText(getParagrahByPostion(data,14,0,37).text) if (isSelected(getParagrahByPostion(data,14,0,37))) else "");fileResult.write(",")
                     
    fileResult.write(replaceText(getParagrahByPostion(data,14,0,43).text
                                 ) +replaceText(getParagrahByPostion(data,14,0,44
                                ).text) + replaceText(getParagrahByPostion(data,14,0,45).text)+ ",")  # 您最急切的需要解决的问题是什么？
    fileResult.write(simpleSelect(getParagrahByPostion(data,16,0,1)) + ",")  # 请问您是否能来参加座谈
    fileResult.write(simpleSelect(getParagrahByPostion(data,18,0,7)) + ",")  # 是否推荐参加座谈会
    fileResult.write(replaceText(getParagrahByPostion(data,18,0,0).text) + ",")
    fileResult.write(filename.split("/")[-1])
    fileResult.write("\n")
    
    open("docx.xml", "a").write(filename+"\n"+ doc._element.xml)

if __name__ == '__main__':

    writeFileFlag = True
    if writeFileFlag:
        open("text.txt", "w").write("清理\n")
        open("docx.xml", "w").write("")
        open("error_file.txt", "w").write("")
        open("result.csv", "w").write("")
        
    fileResult = open("result.csv", "a")
    fileError = open("error_file.txt", "a")
    for root, dirs, files in os.walk(directoryName):
        for fileno, f in enumerate(files):
            filename = os.path.join(root, f)
            print(filename)
            # os.renames(filename,filename.replace(" ",""))
            if f.startswith("."):
                continue
            if testFlag and fileno == 50:
                break

            try:
                parseFile(fileResult, filename)
            except Exception as e:
                fileResult.write("\n")
                fileResult.flush()
                traceback.print_exc()
                fileError.write("error filename : %s\n" % (filename))
        pass
#    pass



